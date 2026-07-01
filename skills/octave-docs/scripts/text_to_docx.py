#!/usr/bin/env python3
"""Convert plain text into a first-pass DOCX using the Octave memo template.

Outline format (markdown-like):
  # Memo Title
  ## Heading
  Body paragraph text.
  Another body paragraph.

Usage:
  python3 text_to_docx.py <outline.txt> <output.docx>
"""
import sys
import zipfile
from pathlib import Path

import docx

ASSET = Path(__file__).parent.parent / "assets" / "Word" / "Octave Memo_US.dotx"

TEMPLATE_CT = b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml"
DOCUMENT_CT = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def load_template(dotx_path: Path, tmp_docx: Path) -> docx.Document:
    """python-docx refuses .dotx directly (content type is 'template', not
    'document'). Patch [Content_Types].xml so it opens as a normal doc."""
    with zipfile.ZipFile(dotx_path) as zin:
        with zipfile.ZipFile(tmp_docx, "w", zipfile.ZIP_DEFLATED) as zout:
            for name in zin.namelist():
                data = zin.read(name)
                if name == "[Content_Types].xml":
                    data = data.replace(TEMPLATE_CT, DOCUMENT_CT)
                zout.writestr(name, data)
    return docx.Document(tmp_docx)


def clear_body(document: docx.Document) -> None:
    """The memo template ships example paragraphs (placeholder body copy,
    closing, signature). Every generated memo must start empty."""
    body = document.element.body
    for p in list(body.findall(f"{W_NS}p")):
        body.remove(p)


def parse_outline(text: str):
    """Returns (title, sections). Each section is (heading, paragraphs)."""
    lines = [ln.rstrip() for ln in text.splitlines() if ln.strip()]
    if not lines or not lines[0].startswith("# "):
        raise ValueError("Outline must start with '# Memo Title'")
    title = lines[0][2:].strip()

    sections = []
    heading = None
    paragraphs = []
    for line in lines[1:]:
        if line.startswith("## "):
            if heading is not None:
                sections.append((heading, paragraphs))
            heading = line[3:].strip()
            paragraphs = []
        else:
            paragraphs.append(line.strip())
    if heading is not None:
        sections.append((heading, paragraphs))

    return title, sections


def build_docx(outline_text: str, output_path: Path) -> None:
    tmp_docx = output_path.with_suffix(".base.docx")
    document = load_template(ASSET, tmp_docx)
    clear_body(document)

    title, sections = parse_outline(outline_text)
    document.add_paragraph(title, style="Heading 1 Memo")

    for heading, paragraphs in sections:
        document.add_paragraph(heading, style="Heading 2")
        for para in paragraphs:
            document.add_paragraph(para, style="Body Text")

    document.save(output_path)
    tmp_docx.unlink(missing_ok=True)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: text_to_docx.py <outline.txt> <output.docx>", file=sys.stderr)
        sys.exit(1)
    outline_path, out_path = Path(sys.argv[1]), Path(sys.argv[2])
    build_docx(outline_path.read_text(), out_path)
    print(f"Wrote {out_path}")
